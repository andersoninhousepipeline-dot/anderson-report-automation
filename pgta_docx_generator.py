"""
DOCX Report Generator for PGT-A Reports
Generates Word documents matching the PDF template with 1:1 precision.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import os
import sys
from io import BytesIO
from datetime import datetime

def set_cell_border(cell, **kwargs):
    """
    Set cell's border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "single", "color": "#0000FF"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existence, if none found, then create one
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('start', 'top', 'end', 'bottom', 'left', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existence, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


class PGTADocxGenerator:
    """Generates DOCX reports for PGT-A with pixel-level precision"""
    
    # Static content (same as PDF template)
    METHODOLOGY_TEXT = """Chromosomal aneuploidy analysis was performed using ChromInst® PGT-A kit from Yikon Genomics (Suzhou) Co., Ltd - China. The Yikon - ChromInst® PGT-A kit with the Genemind - SURFSeq 5000* High-throughput Sequencing Platform allows detection of aneuploidies in all 23 sets of Chromosomes. Probes are not covering the p arm of acrocentric chromosomes as they are rich in repeat regions and RNA markers and devoid of genes. Changes in this region will not be detected. However, these regions have less clinical significance due to the absence of genes. Chromosomal aneuploidy can be detected by copy number variations (CNVs), which represent a class of variation in which segments of the genome have been duplicated (gains) or deleted (losses). Large, genomic copy number imbalances can range from sub-chromosomal regions to entire chromosomes. Inherited and de-novo CNVs (up to 10 Mb) have been associated with many disease conditions. This assay was performed on DNA extracted from embryo biopsy samples."""
    
    MOSAICISM_TEXT = """Mosaicism arises in the embryo due to mitotic errors which lead to the production of karyotypically distinct cell lineages within a single embryo [1]. NGS has the sensitivity to detect mosaicism when 30% or the above cells are abnormal [2]. Mosaicism is reported in our laboratory as follows [3]."""
    
    MOSAICISM_BULLETS = [
        "Embryos with less than 30% mosaicism are considered as euploid.",
        "Embryos with 30% to 50% mosaicism will be reported as low level mosaic, 51% to 80% mosaicism will be reported as high level mosaic.",
        "When three chromosomes or more than three chromosomes showing mosaic change, it will be denoted as complex mosaic.",
        "If greater than 80% mosaicism detected in an embryo it will be considered aneuploid."
    ]
    
    MOSAICISM_CLINICAL = """Clinical significance of transferring mosaic embryos is still under evaluation. Based on Preimplantation Genetic Diagnosis International Society (PGDIS) Position Statement – 2019 transfer of these embryos should be considered only after appropriate counselling of the patient and alternatives have been discussed. Invasive prenatal testing with karyotyping in the amniotic fluid needs to be advised in such cases [4]. As shown in published literature evidence, such transfers can result in normal pregnancy or miscarriage or an offspring with chromosomal mosaicism [5,6,7]."""
    
    LIMITATIONS = [
        "This technique cannot detect point mutations, balanced translocations, inversions, triploidy, uniparental disomy and epigenetic modifications.",
        "Probes used do not cover the p arm of acrocentric chromosomes as they are rich in repeat regions and RNA markers and devoid of genes. Changes in this region will not be detected. However, these regions have less clinical significance due to the absence of genes.",
        "Deletions and duplications with the size of < 10 Mb cannot be detected.",
        "Risk of misinterpretation of the actual embryo karyotype due to the presence of chromosomal mosaicism, either at cleavage-stage or at blastocyst stage may exist.",
        "This technique cannot detect variants of polyploidy and haploidy",
        "NGS without genotyping cannot identify the nature (meiotic or mitotic) nor the parental origin of aneuploidies",
        "Due to the intrinsic nature of chromosomal mosaicism, the chromosomal make-up achieved from a biopsy only may represent a picture of a small part of the embryo and may not necessarily reflect the chromosomal content of the entire embryo. Also, the mosaicism level inferred from a multi-cell TE biopsy might not unequivocally represent the exact chromosomal mosaicism percentage of the TE cells or the inner cell mass constitution."
    ]
    
    REFERENCES = [
        'McCoy, Rajiv C. "Mosaicism in Preimplantation human embryos: when chromosomal abnormalities are the norm." Trends in genetics 33.7 (2017): 448-463.',
        'ESHRE PGT-SR/PGT-A Working Group, et al. "ESHRE PGT Consortium good practice recommendations for the detection of structural and numerical chromosomal aberrations." Human reproduction open 2020.3 (2020): hoaa017.',
        'ESHRE Working Group on Chromosomal Mosaicism, et al. "ESHRE survey results and good practice recommendations on managing chromosomal mosaicism." Hum Reprod Open. 2022 Nov 7;2022(4):hoac044.',
        'Cram, D. S., et al. "PGDIS position statement on the transfer of mosaic embryos 2019." Reproductive biomedicine online 39 (2019): e1-e4.',
        'Victor, Andrea R., et al. "One hundred mosaic embryos transferred prospectively in a single clinic: exploring when and why they result in healthy pregnancies." Fertility and sterility 111.2 (2019): 280-293.',
        'Lin, Pin-Yao, et al. "Clinical outcomes of single mosaic embryo transfer: high-level or low-level mosaic embryo, does it matter?" Journal of clinical medicine 9.6 (2020): 1695.',
        'Kahraman, Semra, et al. "The birth of a baby with mosaicism resulting from a known mosaic embryo transfer: a case report." Human Reproduction 35.3 (2020): 727-733.'
    ]

    def __init__(self, assets_dir="assets/pgta"):
        """Initialize assets and log paths"""
        script_dir = os.path.dirname(os.path.abspath(__file__))
        self.assets_dir = os.path.join(script_dir, assets_dir)
        
        self.header_logo = os.path.join(self.assets_dir, "image_page1_0.png")
        self.footer_banner = os.path.join(self.assets_dir, "image_page1_1.png")
        self.genqa_logo = os.path.join(self.assets_dir, "genqa_logo.png")
        self.signs_image = os.path.join(self.assets_dir, "signs.png")

    # --- OXML PRECISION HELPERS ---
    
    def _set_cell_background(self, cell, fill):
        """Set background shading for a table cell using OXML"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill.replace('#', ''))
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def _set_table_fixed_layout(self, table):
        """Force a table to use a fixed layout so column widths are strictly respected"""
        tbl_pr = table._element.xpath('w:tblPr')[0]
        layout = OxmlElement('w:tblLayout')
        layout.set(qn('w:type'), 'fixed')
        tbl_pr.append(layout)

    def _set_column_widths(self, table, widths_pt):
        """Set exact column widths in points (1 pt = 1/72 inch) for every row/cell"""
        total_width = sum(widths_pt)
        # Set total table width
        tbl_pr = table._element.xpath('w:tblPr')[0]
        tbl_w = OxmlElement('w:tblW')
        tbl_w.set(qn('w:w'), str(int(total_width * 20))) 
        tbl_w.set(qn('w:type'), 'dxa')
        tbl_pr.append(tbl_w)

        # Iterate rows and set each cell's width to ensure parity even if columns[] fails
        for row in table.rows:
            for i, width in enumerate(widths_pt):
                if i < len(row.cells):
                    cell = row.cells[i]
                    tc_pr = cell._tc.get_or_add_tcPr()
                    tc_w = OxmlElement('w:tcW')
                    tc_w.set(qn('w:w'), str(int(width * 20))) 
                    tc_w.set(qn('w:type'), 'dxa')
                    tc_pr.append(tc_w)

    def _set_paragraph_font(self, paragraph, font_name="Segoe UI", font_size=9, bold=False, italic=False, color=None):
        """Apply font styling to every run in a paragraph to ensure 1:1 PDF parity"""
        if not paragraph.runs:
            paragraph.add_run()
        for run in paragraph.runs:
            run.font.name = font_name
            r = run._element
            r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:ascii'), font_name)
            r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:hAnsi'), font_name)
            
            run.font.size = Pt(font_size)
            run.bold = bold
            run.italic = italic
            if color:
                if isinstance(color, str) and color.startswith('#'):
                    run.font.color.rgb = RGBColor.from_string(color[1:])
                else:
                    run.font.color.rgb = color

    def _clean(self, val, default=""):
        """Sanitize values"""
        if val is None: return default
        s = str(val).strip()
        if s.lower() == "nan": return default
        return s

    def _fmt_age(self, val):
        """Format age: '37.0' → '37 Years', '37' → '37 Years', '37 Years' unchanged"""
        import re as _re
        s = self._clean(val)
        if not s:
            return ""
        m = _re.match(r'^(\d+)(?:\.0+)?$', s.strip())
        if m:
            return f"{m.group(1)} Years"
        s = _re.sub(r'^(\d+)\.0+(\s)', r'\1\2', s)
        return s

    # --- GENERATION LOGIC ---

    def generate_docx(self, output_path, patient_data, embryos_data, show_logo=True, show_grid=False):
        """Main entry point for DOCX generation"""
        self.show_grid = show_grid
        doc = Document()
        
        # 1. Page Setup (Margins mirroring PDF exactly)
        sections = doc.sections
        # 1. Page Setup (US Letter: 612pt x 792pt)
        sections = doc.sections
        for section in sections:
            section.page_width = Pt(612)
            section.page_height = Pt(792)
            section.top_margin = Pt(70)
            section.bottom_margin = Pt(60)
            section.left_margin = Pt(58)
            section.right_margin = Pt(58)
            section.header_distance = Pt(20)
            section.footer_distance = Pt(20)
        
        # Global Font Defaults
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(9)
        
        # 2. Cover Page
        self._add_cover_page(doc, patient_data, embryos_data)
        
        # 3. Headers/Footers
        self._setup_page_header_footer(doc, show_logo=show_logo)
        
        # 4. Methodology Page
        doc.add_page_break()
        self._add_methodology_page(doc)
        
        # 5. Embryo Result Pages
        for idx, embryo in enumerate(embryos_data):
            self._add_embryo_page(doc, patient_data, embryo)
            if idx < len(embryos_data) - 1:
                # Page break within _add_embryo_page handles start of new page
                pass
        
        # 6. Save
        doc.save(output_path)
        return output_path

    def _setup_page_header_footer(self, doc, show_logo=True):
        """Setup branding in headers and footers using locked table layouts"""
        for section in doc.sections:
            # Header
            header = section.header
            header.paragraphs[0].clear()
            if show_logo and self.header_logo and os.path.exists(self.header_logo):
                p = header.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(self.header_logo, width=Pt(496))

            # Footer
            footer = section.footer
            footer.paragraphs[0].clear()
            footer_table = footer.add_table(rows=1, cols=2, width=Pt(496))
            self._set_table_fixed_layout(footer_table)
            self._set_column_widths(footer_table, [416, 80])
            
            # Banner
            if show_logo and self.footer_banner and os.path.exists(self.footer_banner):
                c0 = footer_table.rows[0].cells[0]
                p0 = c0.paragraphs[0]
                p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p0.add_run().add_picture(self.footer_banner, width=Pt(416))
            
            # GenQA
            if self.genqa_logo and os.path.exists(self.genqa_logo):
                c1 = footer_table.rows[0].cells[1]
                p1 = c1.paragraphs[0]
                p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p1.add_run().add_picture(self.genqa_logo, width=Pt(65))

    def _add_cover_page(self, doc, patient_data, embryos_data):
        """Cover page mirroring PDF layout and colors"""
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("Preimplantation Genetic Testing for Aneuploidies (PGT-A)")
        self._set_paragraph_font(title, font_name="Calibri", font_size=14, bold=False)
        
        doc.add_paragraph() # Spacer
        
        # Patient Info Table [108, 12, 131, 108, 12, 119] - 6 rows (spouse name combined with patient name)
        # Adjusted widths to give more space to label columns to prevent date wrap
        info_table = doc.add_table(rows=6, cols=6)
        self._apply_grid_to_table(info_table)
        self._set_table_fixed_layout(info_table)
        self._set_column_widths(info_table, [108, 12, 131, 108, 12, 119])
        self._populate_patient_table(info_table, patient_data, is_embryo=False)
        
        doc.add_paragraph() # Spacer
        
        # PNDT Disclaimer
        disclaimer = doc.add_paragraph()
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_paragraph_font(disclaimer, font_name="Calibri", font_size=9.5, italic=True)
        disclaimer.add_run("This test does not reveal sex of the fetus & confers to PNDT act, 1994")
        
        doc.add_paragraph() # Spacer
        
        # Indication
        if 'indication' in patient_data and patient_data['indication']:
            p_ind = doc.add_paragraph()
            self._set_paragraph_font(p_ind, font_name="Calibri", font_size=10, bold=True)
            p_ind.add_run("Indication")
            p_val = doc.add_paragraph(self._clean(patient_data['indication']))
            self._set_paragraph_font(p_val, font_size=9)
            doc.add_paragraph()

        # Results Summary Header
        p_res = doc.add_paragraph()
        self._set_paragraph_font(p_res, font_name="Calibri", font_size=10, bold=True)
        p_res.add_run("Results summary")
        
        # Results Summary Table [50, 95, 185, 80, 86]
        res_table = doc.add_table(rows=len(embryos_data) + 1, cols=5)
        self._apply_grid_to_table(res_table)
        self._set_table_fixed_layout(res_table)
        self._set_column_widths(res_table, [50, 95, 185, 80, 86])
        
        # Row 0: Headers (Peach bg)
        headers = ['S. No.', 'Sample', 'Result', 'MTcopy', 'Interpretation']
        for i, h in enumerate(headers):
            cell = res_table.rows[0].cells[i]
            cell.text = h
            self._set_paragraph_font(cell.paragraphs[0], font_size=9, bold=True)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_cell_background(cell, "F9BE8F")

        # Data rows (F1F1F7 bg)
        for i, emb in enumerate(embryos_data, 1):
            row = res_table.rows[i]
            row.cells[0].text = str(i)
            row.cells[1].text = self._clean(emb.get('embryo_id'))
            
            res_sum = self._clean(emb.get('result_summary'))
            interp = self._clean(emb.get('interpretation'))
            mt = self._clean(emb.get('mtcopy'), 'NA')
            
            # Logic: If any of (autosomes, sex chromosomes, result summary) is abnormal, interpretation is Aneuploid.
            # If all are normal (and not mosaic), interpretation is Euploid.
            auto_val = self._clean(emb.get('autosomes')).upper()
            sex_val = self._clean(emb.get('sex_chromosomes', 'Normal')).upper()
            res_val = res_sum.upper()
            
            # Only auto-derive interpretation when not explicitly set by user
            if not interp:
                if "LOW DNA" in res_val or "INCONCLUSIVE" in res_val:
                    interp = "NA"
                elif self._is_abnormal(auto_val) or self._is_abnormal(sex_val) or self._is_abnormal(res_val):
                    interp = "Aneuploid"
                elif "MOSAIC" in auto_val or "MOSAIC" in sex_val or "MOSAIC" in res_val or "%" in auto_val:
                    interp = self._mosaic_level(auto_val + " " + sex_val)
                else:
                    is_a_n = not auto_val or "NORMAL" in auto_val or "EUPLOID" in auto_val
                    is_s_n = not sex_val or "NORMAL" in sex_val or "EUPLOID" in sex_val
                    is_r_n = not res_val or "NORMAL" in res_val or "EUPLOID" in res_val
                    if is_a_n and is_s_n and is_r_n:
                        interp = "Euploid"

            if "MOSAIC" in interp.upper():
                pass  # keep mosaic percentage
            elif interp.upper() != "EUPLOID":
                mt = "NA"

            row.cells[2].text = res_sum
            row.cells[3].text = mt
            row.cells[4].text = interp
            
            interp_color = self._get_result_color_hex(res_sum, interp)
            
            for c_idx, cell in enumerate(row.cells):
                self._set_cell_background(cell, "F1F1F7")
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Color only Interpretation and Result according to logic
                color = interp_color if c_idx == 4 else None
                self._set_paragraph_font(p, font_size=9, color=color)
        
        # Results Summary Comment (optional, appears below table)
        results_summary_comment = self._clean(patient_data.get('results_summary_comment', ''))
        if results_summary_comment:
            doc.add_paragraph()  # Spacer
            p_comment = doc.add_paragraph(results_summary_comment)
            self._set_paragraph_font(p_comment, font_size=9)

    def _populate_patient_table(self, table, data, is_embryo=False):
        """Standard Patient Info Population Logic"""
        # Patient name and spouse name - spouse on new line
        import re
        patient_name = re.sub(r'\s+', ' ', self._clean(data.get('patient_name'))).strip()
        spouse_name = re.sub(r'\s+', ' ', self._clean(data.get('spouse_name'))).strip()
        # Put spouse on new line if present
        combined_name = f"{patient_name}\n{spouse_name}" if spouse_name else patient_name
        
        rows_map = [
            ("PATIENT NAME", combined_name, "PIN", "pin"),
            ("DATE OF BIRTH/ AGE", "age", "SAMPLE NUMBER", "sample_number"),
            ("REFERRING CLINICIAN", "referring_clinician", "BIOPSY DATE", "biopsy_date"),
            ("HOSPITAL/CLINIC", "hospital_clinic", "SAMPLE COLLECTION DATE", "sample_collection_date"),
            ("SPECIMEN", "specimen", "SAMPLE RECEIPT DATE", "sample_receipt_date"),
            ("BIOPSY PERFORMED BY", "biopsy_performed_by", "REPORT DATE", "report_date")
        ]
        for r_idx, (l1, v1, l2, v2) in enumerate(rows_map):
            if r_idx >= len(table.rows): break
            row = table.rows[r_idx]
            
            # Populate labels and colons
            if l1: row.cells[0].text = l1; row.cells[1].text = ":"
            if l2: row.cells[3].text = l2; row.cells[4].text = ":"
            
            # Populate cleaned values - first row has combined name directly
            if v1:
                if r_idx == 0:  # First row - combined name already a string
                    row.cells[2].text = v1
                elif v1 == "age":
                    row.cells[2].text = self._fmt_age(data.get(v1))
                else:
                    row.cells[2].text = self._clean(data.get(v1))
            if v2: row.cells[5].text = self._clean(data.get(v2))
            
            for cell_idx, cell in enumerate(row.cells):
                self._set_cell_background(cell, "F1F1F7")
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                self._set_paragraph_font(cell.paragraphs[0], font_name="Segoe UI", font_size=10, bold=True)
                p_fmt = cell.paragraphs[0].paragraph_format
                p_fmt.space_before = Pt(2)
                p_fmt.space_after = Pt(2)
                
                # Set cell alignment to match PDF strictly left aligned for values
                # Fixed: Use cell_idx from enumerate instead of row.cells.index(cell) to avoid tuple.index error
                if cell_idx in [0, 3]:  # Label columns
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                    
                    # Logic: PIN label in embryo banner should be right-aligned (flushed to colon)
                    # Page 1 and other labels should remain left-aligned with 12pt padding
                    label_text = cell.text.strip()
                    if is_embryo and label_text == "PIN":
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        cell.paragraphs[0].paragraph_format.left_indent = Pt(0)
                        cell.paragraphs[0].paragraph_format.right_indent = Pt(12)
                    else:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        cell.paragraphs[0].paragraph_format.left_indent = Pt(4)
                elif cell_idx in [1, 4]:  # Colon columns
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:  # Value columns
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _add_methodology_page(self, doc):
        """Methods, Limitations, and References with natural flow but orphan protection"""
        # Content sections
        sections = [
            ("Methodology", self.METHODOLOGY_TEXT, None),
            ("Conditions for reporting mosaicism", self.MOSAICISM_TEXT, self.MOSAICISM_BULLETS),
            (None, self.MOSAICISM_CLINICAL, None),
            ("Limitations", None, self.LIMITATIONS),
            ("References", None, [f"{i}. {r}" for i, r in enumerate(self.REFERENCES, 1)])
        ]
        
        for head, body, bullets in sections:
            if head:
                p = doc.add_paragraph()
                self._set_paragraph_font(p, font_size=11, bold=True)
                p.add_run(head)
                # Only keep with next if there's content following
                if body or bullets:
                    p.paragraph_format.keep_with_next = True
            
            if body:
                p = doc.add_paragraph(body)
                self._set_paragraph_font(p, font_size=9)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # Only keep with bullets if they exist
                if bullets:
                    p.paragraph_format.keep_with_next = True
                
            if bullets:
                for i, b in enumerate(bullets):
                    p = doc.add_paragraph(b, style='List Bullet')
                    self._set_paragraph_font(p, font_size=9)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    # Keep bullets together naturally
                    if i < len(bullets) - 1:
                        p.paragraph_format.keep_with_next = True
            doc.add_paragraph()

    def _add_embryo_page(self, doc, patient_data, embryo_data):
        """Individual Embryo Result Page with exact PDF metrics"""
        doc.add_page_break()
        
        # 1. Banner [Total: 490pt] - Match exact cover page positioning
        banner = doc.add_table(rows=2, cols=6)
        self._apply_grid_to_table(banner)
        self._set_table_fixed_layout(banner)
        # Optimized layout: Push PIN block right. PATIENT NAME (82), Colons (12x2), PIN label (24).
        self._set_column_widths(banner, [82, 12, 242, 24, 12, 118])
        # Ensure table is aligned to the left like cover page
        banner.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._populate_patient_table(banner, patient_data, is_embryo=True)

        doc.add_paragraph()
        
        # 2. Embryo ID - Use embryo_id_detail for detail pages, fallback to embryo_id
        eid = self._clean(embryo_data.get('embryo_id_detail')) or self._clean(embryo_data.get('embryo_id'))
        p_eid = doc.add_paragraph()
        self._set_paragraph_font(p_eid, font_name="Calibri", font_size=12, bold=True, color="#1F497D")
        p_eid.add_run(f"EMBRYO: {eid}")
        
        # 3. Summary [Total: 490pt]
        res = self._clean(embryo_data.get('result_summary'))
        interp = self._clean(embryo_data.get('interpretation'))
        auto = self._clean(embryo_data.get('autosomes'))
        sex = self._clean(embryo_data.get('sex_chromosomes'))
        
        # Logic: If any of (autosomes, sex chromosomes, result summary) is abnormal, interpretation is Aneuploid.
        # If all are normal (and not mosaic), interpretation is Euploid.
        auto_val = auto.upper()
        sex_val = sex.upper()
        res_val = res.upper()
        # Only auto-derive interpretation when not explicitly set by user
        if not interp:
            if self._is_abnormal(auto_val) or self._is_abnormal(sex_val) or self._is_abnormal(res_val):
                interp = "Aneuploid"
            elif "MOSAIC" in auto_val or "MOSAIC" in sex_val or "MOSAIC" in res_val or "%" in auto_val:
                interp = self._mosaic_level(auto_val + " " + sex_val)
            else:
                is_a_n = not auto_val or "NORMAL" in auto_val or "EUPLOID" in auto_val
                is_s_n = not sex_val or "NORMAL" in sex_val or "EUPLOID" in sex_val
                is_r_n = not res_val or "NORMAL" in res_val or "EUPLOID" in res_val
                if is_a_n and is_s_n and is_r_n:
                    interp = "Euploid"

        mt = self._clean(embryo_data.get('mtcopy'), 'NA')
        if "MOSAIC" in interp.upper():
            if not mt or mt.upper() == "NA":
                import re as _re
                pcts = _re.findall(r'(\d+)%', auto_val)
                if pcts:
                    mt = f"{max(int(p) for p in pcts)}%"
        elif interp.upper() != "EUPLOID":
            mt = "NA"
        
        result_desc_text = self._clean(embryo_data.get('result_description', ''))
        interp_color = self._get_result_color_hex(res, interp)
        # Force black color ONLY for the "Result:" description row in details section
        res_row_color = "#000000"
        details = [
            ("Result:", res, res_row_color),
            ("Autosomes:", auto, self._get_status_color_docx(auto)),
            ("Sex Chromosomes:", sex, "#0000FF" if "MOSAIC" in sex.upper() else ("#FF0000" if "ABNORMAL" in sex.upper() else "#000000")),
            ("Interpretation:", interp, interp_color),
            ("MTcopy:", mt, "#000000")
        ]
        
        d_table = doc.add_table(rows=len(details), cols=1)
        self._apply_grid_to_table(d_table)
        self._set_table_fixed_layout(d_table)
        self._set_column_widths(d_table, [490])
        for idx, (label, val, color) in enumerate(details):
            cell = d_table.rows[idx].cells[0]
            self._set_cell_background(cell, "F1F1F7")
            p = cell.paragraphs[0]
            self._set_paragraph_font(p, font_size=9, bold=True)
            p.add_run(f"{label} ")
            run_val = p.add_run(val)
            self._set_paragraph_font(p, font_size=9, bold=False, color=color)
            p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)

        doc.add_paragraph()
        
        # 4. Chart
        p_ch = doc.add_paragraph()
        self._set_paragraph_font(p_ch, font_size=10, bold=True)
        p_ch.add_run("COPY NUMBER CHART")
        if embryo_data.get('cnv_image_path') and os.path.exists(embryo_data['cnv_image_path']):
            doc.add_picture(embryo_data['cnv_image_path'], width=Pt(496))
        
        doc.add_paragraph()
        
        # 5. CNV Status Table [Total: 496pt] - Skip for Inconclusive results
        result_summary = self._clean(embryo_data.get('result_summary', ''))
        result_desc = self._clean(embryo_data.get('result_description', ''))
        is_inconclusive = "INCONCLUSIVE" in result_summary.upper() or "INCONCLUSIVE" in result_desc.upper() or "INCONCLUSIVE" in interp.upper()
        
        # Add inconclusive comment under CNV chart if present
        if is_inconclusive:
            inconclusive_comment = self._clean(embryo_data.get('inconclusive_comment', ''))
            if inconclusive_comment:
                comment_p = doc.add_paragraph(inconclusive_comment)
                self._set_paragraph_font(comment_p, font_size=11)
        if not is_inconclusive:
            chr_statuses = embryo_data.get('chromosome_statuses', {})
            mosaic_map = embryo_data.get('mosaic_percentages', {})
            
            autosomes = str(embryo_data.get('autosomes', '')).upper()
            sex_chrs = str(embryo_data.get('sex_chromosomes', '')).upper()
            
            # Mosaic status codes that require a Mosaic(%) row
            _MOSAIC_CODES = {'M', 'MG', 'ML', 'SMG', 'SML'}
            
            import re as re_mos
            # Check if any chromosome has a mosaic CNV status code
            has_mosaic_status = any(
                str(v).strip().upper() in _MOSAIC_CODES
                for v in chr_statuses.values()
            )
            # Check if any mosaic percentage has a real numeric value
            has_mosaic_pct = any(
                v and str(v).strip() and str(v).strip() != '-' and re_mos.search(r'\d', str(v))
                for v in mosaic_map.values()
            )
            has_mosaic = has_mosaic_status or has_mosaic_pct
            
            is_autosomes_normal = 'NORMAL' in autosomes or 'EUPLOID' in autosomes or not autosomes.strip()
            is_sex_mosaic = 'MOSAIC' in sex_chrs
            
            if is_autosomes_normal and is_sex_mosaic:
                has_mosaic = False
                
            cnv_fs = 7  # single value controls both Chromosome and CNV status rows
            num_rows = 3 if has_mosaic else 2
            cnv_table = doc.add_table(rows=num_rows, cols=23)
            self._apply_grid_to_table(cnv_table)
            self._set_table_fixed_layout(cnv_table)
            self._set_column_widths(cnv_table, [75] + [19.13]*22)

            # Header Row
            cnv_table.rows[0].cells[0].text = "Chromosome"
            for i in range(1, 23): cnv_table.rows[0].cells[i].text = str(i)

            # Status Row
            cnv_table.rows[1].cells[0].text = "CNV status"
            for i in range(1, 23):
                cell = cnv_table.rows[1].cells[i]
                stat = str(chr_statuses.get(str(i), 'N'))
                color = self._get_status_color_docx(stat)
                p = cell.paragraphs[0]
                if '/' in stat:
                    parts = stat.split('/', 1)
                    run1 = p.add_run(parts[0] + '/')
                    run1.add_break()          # <w:br/> forces wrap after slash
                    p.add_run(parts[1])
                else:
                    cell.text = stat
                self._set_paragraph_font(p, font_size=cnv_fs, bold=True, color=color)

            # Mosaic Row - percentage values colored by their chromosome's status
            if has_mosaic:
                # Label cell
                cnv_table.rows[2].cells[0].text = "Mosaic (%)"
                self._set_paragraph_font(cnv_table.rows[2].cells[0].paragraphs[0], font_size=cnv_fs, bold=True)
                for i in range(1, 23):
                    perc = str(mosaic_map.get(str(i), '-'))
                    if not perc.strip():
                        perc = '-'
                    cnv_table.rows[2].cells[i].text = perc
                    # Color the percentage value to match the chromosome's CNV status color
                    stat = str(chr_statuses.get(str(i), 'N'))
                    perc_color = self._get_status_color_docx(stat)
                    self._set_paragraph_font(cnv_table.rows[2].cells[i].paragraphs[0], font_size=cnv_fs, bold=True, color=perc_color)

            for row in cnv_table.rows:
                for c_idx, cell in enumerate(row.cells):
                    self._set_cell_background(cell, "F1F1F7")
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if c_idx == 0: p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # Skip re-formatting mosaic row cells (row index 2) since they already have per-cell colors
                    if has_mosaic and row == cnv_table.rows[2]:
                        continue
                    self._set_paragraph_font(p, font_size=cnv_fs, bold=True)

        doc.add_paragraph()
        self._add_signature_section(doc)

    def _add_signature_section(self, doc):
        """Pixel-Perfect 3-Column Signature Section"""
        table = doc.add_table(rows=2, cols=3)
        self._set_table_fixed_layout(table)
        self._set_column_widths(table, [156, 156, 156])
        
        if self.signs_image and os.path.exists(self.signs_image):
            p = table.rows[0].cells[1].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(self.signs_image, width=Pt(300))
        
        sigs = [
            ("Dr. Meena G", "Associate Consultant"),
            ("Dr. Manju R", "Medical Geneticist"),
            ("Dr. Shivani P", "Managing Director")
        ]
        for i, (name, title) in enumerate(sigs):
            cell = table.rows[1].cells[i]
            p1 = cell.paragraphs[0]; p1.text = name; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_paragraph_font(p1, font_size=11)
            p2 = cell.add_paragraph(title); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_paragraph_font(p2, font_size=11)

    def _mosaic_level(self, combined_text):
        """Determine Low/High/Complex mosaic level from text containing mosaic percentages."""
        import re as _re
        pcts = [int(p) for p in _re.findall(r'(\d+)%', combined_text)]
        entries = len(_re.findall(r'(?:mos|\(\~?\d+%)', combined_text, _re.IGNORECASE))
        if entries >= 3 or len(pcts) >= 3:
            return "Complex mosaic"
        if pcts:
            return "High level mosaic" if max(pcts) >= 51 else "Low level mosaic"
        return "Low level mosaic"

    def _is_abnormal(self, val):
        """Helper to determine if a result field indicates an abnormality (not Normal/Euploid/Mosaic)"""
        if not val: return False
        v = str(val).upper()
        if "ABNORMAL" in v: return True
        if any(x in v for x in ["NORMAL", "EUPLOID", "NO COPY NUMBER ABNORMALITY"]):
            return False
        if "MOSAIC" in v or "%" in v:  # % indicates mosaic percentage (e.g. +21(~54%))
            return False
        if v == "NA": return False
        return True

    def _get_result_color_hex(self, res, interp):
        """Standard Results Color Map - Euploid=black, Aneuploid=red, Mosaic=blue"""
        i = str(interp).upper()
        # Euploid = Black (check first for explicit euploid)
        if "EUPLOID" in i and "ANEUPLOID" not in i:
            return "#000000"
        if any(k in i for k in ["ANEUPLOID", "ABNORMAL"]) or i.strip() == "(-)": return "#FF0000"
        r = str(res).upper()
        if any(k in r for k in ["ANEUPLOID", "ABNORMAL"]): return "#FF0000"
        # Blue for any mosaic interpretation or result
        if "MOSAIC" in i or "MOSAIC" in r:
            return "#0000FF"
        return "#000000"

    def _get_status_color_docx(self, status):
        """CNV Status Color Map for Autosomes
        Blue (mosaic) = Has % sign (e.g., +15(~30%), -20(~51%), dup(9)...(~32%))
        Red (non-mosaic) = del/dup/-/+ without %, or CNV status L/G/SL/SG
        Black = Normal/Euploid
        """
        s = str(status).upper()
        original = str(status)
        
        # Check for Normal/Euploid first
        if 'NORMAL' in s or 'EUPLOID' in s or not original.strip():
            return "#000000"

        # Specific phrase: Multiple Mosaic Chromosome complement → blue
        if 'MULTIPLE MOSAIC CHROMOSOME COMPLEMENT' in s:
            return "#0000FF"

        # Mosaic = has % sign
        if '%' in original:
            return "#0000FF"
        
        # Non-mosaic abnormalities (no % sign)
        if any(x in s for x in ['DEL(', 'DUP(', 'STATUS L', 'STATUS G', 'STATUS SL', 'STATUS SG', ' SL', ' SG', ' L,', ' G,', ' L ', ' G ']):
            return "#FF0000"
        if s.endswith(' L') or s.endswith(' G'):
            return "#FF0000"
        # Check for +/- patterns (e.g., -16, +7, -22)
        import re
        if re.search(r'^[+-]\d+', original) or re.search(r',[+-]?\d+$', original):
            return "#FF0000"
        if 'CNV STATUS' in s:
            return "#FF0000"
        
        # Check for simple abbreviations at word boundaries
        words = s.split()
        for word in words:
            if word in ['MULTIPLE', 'CHROMOSOMAL', 'ABNORMALITIES']: # Handled elsewhere usually but catching edge cases
                if "MULTIPLE CHROMOSOMAL ABNORMALITIES" in s:
                    return "#FF0000"
            for abbrev in ['L', 'G', 'SL', 'SG', 'SL/SG', 'SG/SL']:
                # Need to be careful with exact matches or substrings like SL/SG
                if abbrev in s.split() or abbrev in s.split(','):
                    return "#FF0000"
            for abbrev in ['MG', 'ML', 'SMG', 'SML', 'M', 'SML/SMG', 'SMG/SML']:
                if abbrev in s.split() or abbrev in s.split(','):
                    return "#0000FF"
                    
        # Explicit checks for the new combinations
        if any(x in s for x in ["SL/SG", "SG/SL", "MULTIPLE CHROMOSOMAL ABNORMALITIES"]):
            return "#FF0000"
        if any(x in s for x in ["SML/SMG", "SMG/SML"]):
            return "#0000FF"
            
        return "#000000"

    def _apply_grid_to_table(self, table):
        """Apply lite white grid lines to table if enabled"""
        if not hasattr(self, 'show_grid') or not self.show_grid:
            return
            
        grid_color = "E0E0E0" # Lite white/grey
        for row in table.rows:
            for cell in row.cells:
                set_cell_border(
                    cell,
                    top={"sz": 4, "val": "single", "color": grid_color},
                    bottom={"sz": 4, "val": "single", "color": grid_color},
                    start={"sz": 4, "val": "single", "color": grid_color},
                    end={"sz": 4, "val": "single", "color": grid_color}
                )
